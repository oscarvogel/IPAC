from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
OUT = BASE / "Propuesta_Comercial_IPAC_Vogel_Consultoria.docx"
LOGO = BASE / "assets" / "logo-vogel.png"
LOGO_DARK = BASE / "assets" / "logo-vogel-dark.png"

BODY_FONT = "DM Sans"
HEAD_FONT = "Syne"

DEEP = RGBColor(15, 42, 68)
NAVY = RGBColor(11, 32, 53)
SLATE = RGBColor(22, 47, 73)
BLUE = RGBColor(30, 95, 168)
BRIGHT = RGBColor(25, 110, 207)
BLUE_LIGHT = RGBColor(139, 197, 255)
AMBER = RGBColor(242, 169, 0)
LIGHT_BLUE = "EAF3FF"
LIGHT_GRAY = "F8FAFC"
PALE_AMBER = "FFF7E0"
MID_GRAY = RGBColor(84, 99, 118)
MUTED = RGBColor(142, 168, 195)
BLACK = RGBColor(25, 29, 35)
WHITE = RGBColor(255, 255, 255)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D7DEE8", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_run(run, size=10.5, color=BLACK, bold=False, italic=False, font=BODY_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_para(doc, text="", size=10.5, color=BLACK, bold=False, italic=False, align=None, before=0, after=6, line=1.12, font=BODY_FONT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run(r, size=size, color=color, bold=bold, italic=italic, font=font)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run(r, size=16.5 if level == 1 else 13, color=DEEP if level == 1 else BLUE, bold=True, font=HEAD_FONT)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run(r, size=10.1, color=BLACK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run(r, size=10.1, color=BLACK)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE, border="D7E8FA", title_color=DEEP, spacer=True):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_border(cell, color=border, size="8")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run(r, size=10.2, color=title_color, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run(r2, size=9.8, color=BLACK)
    if spacer:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def set_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.14


def add_header_footer(section):
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO.exists():
        run = hp.add_run()
        run.add_picture(str(LOGO), width=Cm(0.95))
    r = hp.add_run("  Vogel Consultoría")
    set_run(r, size=8.8, color=DEEP, bold=True, font=HEAD_FONT)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("vogelconsultoria.com.ar  |  oscarvogel@gmail.com  |  3743 667526")
    set_run(r1, size=8.2, color=MID_GRAY)


def add_plan_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Plan", "Abono mensual", "Horas incluidas", "Uso recomendado"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_fill(cell, "0F2A44")
        set_cell_border(cell, color="0F2A44")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=9.4, color=WHITE, bold=True)

    rows = [
        ("Inicial", "$250.000", "Hasta 10 h/mes", "Avance gradual y soporte básico."),
        ("Recomendado", "$350.000", "Hasta 15 h/mes", "Equilibrio entre avance mensual, soporte y costo."),
        ("Avanzado", "$500.000", "Hasta 25 h/mes", "Mayor velocidad de implementación y ajustes."),
    ]
    for row_idx, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_fill(cells[i], "FFFFFF" if row_idx != 2 else "EAF3FF")
            set_cell_border(cells[i], color="D7DEE8")
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run(r, size=9.1, color=BLACK, bold=(row_idx == 2 and i in (0, 1)))
    set_table_width(table, [Inches(1.25), Inches(1.35), Inches(1.35), Inches(2.55)])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_roadmap_table(doc):
    table = doc.add_table(rows=1, cols=3)
    headers = ["Mes", "Entregables principales", "Beneficio para IPAC"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_fill(cell, "0F2A44")
        set_cell_border(cell, color="0F2A44")
        set_cell_margins(cell)
        r = cell.paragraphs[0].add_run(h)
        set_run(r, size=9.1, color=WHITE, bold=True)
    rows = [
        ("1", "Servidor, usuarios, sucursales, alumnos, carreras/cursos y conceptos base.", "Primer acceso web ordenado desde Posadas y Eldorado."),
        ("2", "Matrícula, generación de cuotas, descuentos por convenio y recargos.", "Menos carga manual y deuda más visible."),
        ("3", "Pagos, pagos a cuenta, estado de cuenta del alumno y reportes de deuda.", "Mejor control de cobranzas y consultas más rápidas."),
        ("4", "Caja por usuario/sucursal, ingresos, egresos, retiros, pases y cierre.", "Control interno diario y menos dependencia de una sola máquina."),
        ("5", "Recibos, exportaciones Excel/PDF, migración inicial y ajustes con datos reales.", "Operación más completa y reportes utilizables por administración."),
        ("6+", "ARCA, Mercado Pago, conciliación, tablero o mejoras priorizadas.", "Evolución según uso real y prioridades de dirección."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_fill(cells[i], "FFFFFF")
            set_cell_border(cells[i], color="D7DEE8")
            set_cell_margins(cells[i], top=115, bottom=115)
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run(r, size=8.55, color=BLACK, bold=(i == 0))
    set_table_width(table, [Inches(0.55), Inches(3.25), Inches(2.7)])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build():
    doc = Document()
    set_document_styles(doc)
    for section in doc.sections:
        add_header_footer(section)

    # Cover page.
    add_para(doc, "PROPUESTA COMERCIAL", size=10.2, color=AMBER, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=14, after=8)
    if LOGO_DARK.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_DARK), width=Cm(5.3))
        p.paragraph_format.space_after = Pt(18)
    elif LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Cm(4.0))
        p.paragraph_format.space_after = Pt(18)
    title = add_para(doc, "Sistema web de administración, tesorería y cobranzas", size=23, color=DEEP, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, line=1.0, font=HEAD_FONT)
    subtitle = add_para(doc, "IPAC Posadas y Eldorado", size=14.5, color=BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=22, font=HEAD_FONT)
    add_callout(
        doc,
        "Objetivo de la propuesta",
        "Implementar un sistema web progresivo que permita trabajar desde ambas sucursales, ordenar la administración, mejorar el control de cuotas y caja, y evolucionar por etapas bajo un abono mensual sostenible.",
        fill=LIGHT_BLUE,
        border="F2A900",
    )
    add_para(doc, "Preparado por Vogel Consultoría", size=10.8, color=DEEP, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=2, font=HEAD_FONT)
    add_para(doc, "Sistemas a medida, dashboards, automatización e inteligencia artificial para empresas.", size=9.5, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "Junio 2026", size=9.5, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    doc.add_page_break()

    add_heading(doc, "1. Resumen ejecutivo")
    add_para(
        doc,
        "IPAC necesita dejar atrás una operación administrativa condicionada por el uso local del sistema y avanzar hacia una plataforma web segura, accesible desde Posadas y Eldorado, con información centralizada y control por sucursal.",
    )
    add_para(
        doc,
        "La propuesta se plantea como un servicio mensual evolutivo: en lugar de congelar un alcance enorme desde el primer día, se construyen módulos priorizados, se entregan avances mensuales y se ajusta el sistema según el uso real de la institución.",
    )
    add_callout(
        doc,
        "Recomendación comercial",
        "Contratar una puesta en marcha inicial y luego un abono mensual recomendado de $350.000, con servidor incluido y hasta 15 horas mensuales para desarrollo, soporte y mejoras evolutivas.",
        fill=PALE_AMBER,
        border="F2A900",
    )

    add_heading(doc, "2. Situación actual detectada")
    for item in [
        "La operación depende de equipos locales y no cuenta con una lógica web centralizada.",
        "Las dos sucursales trabajan con información separada o sin intercambio fluido.",
        "El volumen actual aproximado es de 220 alumnos entre Posadas y Eldorado.",
        "Procesos críticos como cuotas, caja, pagos, reportes y actualizaciones dependen de tareas manuales.",
        "El crecimiento futuro requiere ordenar roles, permisos, auditoría, backups y acceso seguro.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Beneficios concretos para IPAC")
    benefits = [
        ("Acceso desde ambas sucursales", "Posadas y Eldorado podrán operar sin depender de una única máquina física."),
        ("Información centralizada", "Alumnos, cuotas, pagos y movimientos estarán disponibles en una única base controlada."),
        ("Mejor control de deuda", "La administración podrá consultar morosos, pagos a cuenta, cuotas vencidas y estado de cada alumno."),
        ("Caja más ordenada", "Cada usuario y sucursal podrá registrar movimientos, cierres, arqueos y reportes diarios."),
        ("Menos trabajo manual", "Recargos, descuentos por convenio, reportes y exportaciones se automatizan progresivamente."),
        ("Crecimiento por etapas", "ARCA, Mercado Pago, portal alumno y gestión académica se incorporan sin frenar la primera versión."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Beneficio", "Impacto esperado"]):
        cell = table.cell(0, i)
        set_cell_fill(cell, "0F2A44")
        set_cell_border(cell, color="0F2A44")
        set_cell_margins(cell)
        r = cell.paragraphs[0].add_run(h)
        set_run(r, size=9.3, color=WHITE, bold=True)
    for label, detail in benefits:
        cells = table.add_row().cells
        for i, value in enumerate((label, detail)):
            set_cell_fill(cells[i], "FFFFFF")
            set_cell_border(cells[i], color="D7DEE8")
            set_cell_margins(cells[i], top=115, bottom=115)
            r = cells[i].paragraphs[0].add_run(value)
            set_run(r, size=9.0, color=BLACK, bold=(i == 0))
    set_table_width(table, [Inches(2.05), Inches(4.45)])

    doc.add_page_break()

    add_heading(doc, "4. Alcance inicial recomendado")
    add_para(doc, "La primera versión se enfoca en administración, tesorería y cobranzas. La gestión pedagógica queda planificada como etapa futura para no demorar la solución operativa principal.")
    for item in [
        "Acceso web seguro con usuarios y contraseñas.",
        "Sucursales Posadas y Eldorado, con permisos y visibilidad por sede.",
        "Gestión de alumnos, matrícula, conceptos, carreras/cursos y datos de contacto.",
        "Generación de cuotas por ciclo, descuentos por convenio y recargos automáticos.",
        "Registro de pagos, pagos a cuenta, saldos y estado de cuenta del alumno.",
        "Caja por usuario y sucursal, ingresos, egresos, retiros, pases, cierres y arqueos.",
        "Recibos, exportaciones a Excel/PDF y reportes básicos de deuda, caja y pagos.",
        "Base preparada para integrar factura ARCA, Mercado Pago y portal de alumnos.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. Modalidad de trabajo")
    for item in [
        "Cada mes se acuerdan prioridades y entregables según la bolsa de horas incluida.",
        "Las tareas no abordadas en el mes quedan priorizadas para el mes siguiente o se cotizan como adicional.",
        "Las urgencias sobre módulos ya entregados tienen prioridad frente a nuevas funcionalidades.",
        "El cliente participa validando datos, reglas, comprobantes, reportes y pruebas de operación.",
    ]:
        add_number(doc, item)

    add_heading(doc, "6. Roadmap tentativo de implementación")
    add_roadmap_table(doc)
    add_para(doc, "Los plazos pueden ajustarse según la disponibilidad de datos reales, definiciones del equipo IPAC y complejidad de integraciones externas.", size=9.5, color=MID_GRAY, italic=True)

    add_heading(doc, "7. Inversión y abono mensual")
    add_para(doc, "Se propone una contratación con puesta en marcha inicial y luego abono mensual. Este modelo permite avanzar de manera constante sin exigir una inversión inicial de sistema completo cerrado.")
    add_callout(doc, "Puesta en marcha inicial", "$400.000 por única vez. Incluye preparación del entorno, servidor, base técnica, estructura inicial del sistema, backups, configuración de acceso y planificación del backlog.", fill=LIGHT_GRAY, border="F2A900")
    add_plan_table(doc)
    add_callout(
        doc,
        "Plan recomendado",
        "Para IPAC se recomienda el plan de $350.000 mensuales, con hasta 15 horas por mes. Es el punto más equilibrado para avanzar sin sobredimensionar el costo.",
        fill=LIGHT_BLUE,
        border="8BC5FF",
    )

    add_heading(doc, "8. Servidor incluido")
    add_para(doc, "El abono incluye alojamiento en servidor cloud administrado, certificado SSL, mantenimiento técnico básico y copias de seguridad periódicas.")
    for item in [
        "Recursos iniciales equivalentes a 4 vCPU, 8 GB RAM, 30 GB NVMe y transferencia suficiente para el uso institucional previsto.",
        "El valor contempla la tarifa promocional vigente del proveedor durante el primer año; luego, una variación significativa podrá trasladarse previa comunicación.",
        "Si el crecimiento requiere más almacenamiento, tráfico o capacidad, se informará antes de modificar el costo.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    add_heading(doc, "9. Alcance sujeto a validación")
    add_para(doc, "Los siguientes puntos forman parte del roadmap, pero requieren materiales y validación técnica antes de comprometer fecha cerrada:")
    for item in [
        "Migración completa desde Excel: requiere exportaciones reales o anonimizadas de alumnos, cuotas, pagos, conceptos y saldos.",
        "Factura ARCA: requiere CUIT, punto de venta, certificados, condición fiscal, ejemplos de comprobantes y definición de permisos.",
        "Mercado Pago y QR: requiere credenciales, modalidad de cobro, política de comisiones y definición de conciliación.",
        "Portal del alumno, responsable y gestión pedagógica/académica: recomendados para etapas futuras, luego de estabilizar administración y cobranzas.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "10. Materiales necesarios para iniciar")
    for item in [
        "Capturas del sistema actual.",
        "Ejemplo de ficha de alumno.",
        "Ejemplo de cuota, estado de cuenta y recibo.",
        "Ejemplo de cierre de caja y reportes actuales.",
        "Exportación Excel anonimizada de alumnos, cuotas, pagos y conceptos.",
        "Listado de usuarios por sucursal y permisos esperados.",
        "Datos fiscales y certificados necesarios para ARCA.",
        "Información operativa de Mercado Pago, QR o Posnet utilizado actualmente.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11. Continuidad luego del período mínimo")
    add_para(
        doc,
        "El abono evolutivo tiene una permanencia mínima recomendada de 6 meses para completar una primera versión operativa y estabilizar el uso. Finalizado ese período, IPAC podrá continuar con el abono de desarrollo y soporte, o pasar a un plan reducido de continuidad operativa para mantener el sistema activo.",
    )
    add_callout(
        doc,
        "Plan de continuidad operativa",
        "$120.000 mensuales. Incluye servidor cloud, certificado SSL, backups, monitoreo básico, mantenimiento técnico mínimo, correcciones críticas menores y soporte operativo limitado. No incluye nuevas funcionalidades, integraciones nuevas, reportes nuevos, soporte intensivo ni urgencias fuera de horario.",
        fill=LIGHT_GRAY,
        border="F2A900",
    )
    for item in [
        "Si el cliente no continúa con el abono evolutivo ni contrata continuidad operativa, se coordinará la exportación de la información disponible y la baja ordenada del servicio de hosting.",
        "La baja del servicio deberá solicitarse con 30 días de anticipación.",
        "La infraestructura permanecerá activa únicamente mientras se encuentre abonado el plan vigente o el plan de continuidad operativa.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    add_heading(doc, "12. Condiciones comerciales y actualización")
    for item in [
        "El abono mensual incluye la bolsa de horas indicada, servidor, mantenimiento técnico básico, backups y soporte razonable.",
        "Las horas no utilizadas no se acumulan, salvo acuerdo previo por escrito.",
        "Los requerimientos nuevos se priorizan dentro del backlog mensual.",
        "Cambios grandes, urgencias fuera de horario, integraciones externas complejas o tareas que excedan la bolsa mensual podrán cotizarse aparte.",
        "Los valores expresados en pesos argentinos tendrán actualización trimestral conforme a la variación acumulada del IPC publicado por INDEC.",
        "La primera actualización se aplicará al cumplirse tres meses desde el inicio del servicio.",
        "La puesta en marcha mantiene su valor por 15 días desde la emisión de la propuesta; vencido ese plazo, podrá actualizarse.",
        "Costos extraordinarios de infraestructura, licencias o servicios de terceros no contemplados originalmente podrán trasladarse al abono previa comunicación y justificación.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "13. Cierre")
    add_para(doc, "La propuesta busca que IPAC gane control operativo sin esperar un desarrollo enorme antes de usar el sistema. La institución obtiene una plataforma viva, accesible desde ambas sucursales, con entregas mensuales y capacidad de crecer hacia facturación, pagos online, autogestión y gestión académica.")
    add_callout(
        doc,
        "Próximo paso sugerido",
        "Aprobar la puesta en marcha, confirmar el plan mensual recomendado y compartir los materiales iniciales para comenzar con servidor, usuarios, sucursales y estructura de alumnos.",
        fill=PALE_AMBER,
        border="F2A900",
        spacer=False,
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(path)
